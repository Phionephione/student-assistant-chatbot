from flask import Flask, render_template, request, redirect, url_for, session, jsonify, send_file
import google.generativeai as genai
import os
from dotenv import load_dotenv
from docx import Document
from io import BytesIO

load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "a_default_secret_key_for_development") # Use a strong secret key

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

try:
    genai.configure(api_key=GOOGLE_API_KEY)
    model = genai.GenerativeModel('gemini-1.5-pro') # Use the correct model name for your key
    print("Gemini API Initialized Successfully.")
except Exception as e:
    print(f"FATAL ERROR: Could not initialize Gemini API: {e}")
    model = None

# In-memory storage (for demonstration only)
users = {}
conversation_history = {}


@app.route("/")
def index():
    if 'username' in session:
        return redirect(url_for('chat'))
    return redirect(url_for('login'))


@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form.get("username")
        email = request.form.get("email")
        password = request.form.get("password")
        if username in users:
            return render_template("register.html", error="Username already exists")
        users[username] = {"email": email, "password": password}
        conversation_history[username] = []
        return redirect(url_for("login"))
    return render_template("register.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")
        if username in users and users[username]["password"] == password:
            session['username'] = username
            return redirect(url_for("chat"))
        return render_template("login.html", error="Invalid username or password")
    return render_template("login.html")


@app.route("/chat")
def chat():
    if 'username' in session:
        return render_template("chat.html", name=session['username'])
    return redirect(url_for('login'))


@app.route("/logout")
def logout():
    session.pop('username', None)
    return redirect(url_for("login"))


@app.route("/get_response", methods=["POST"])
def get_response():
    if model is None:
        return jsonify({"response": "Error: The AI model is not initialized. Please check the server logs."})

    user_input = request.form["user_input"]
    username = session['username']

    try:
        gemini_response = model.generate_content(user_input)
        answer_text = gemini_response.text

        conversation_history[username].append({
            "question": user_input,
            "answer": answer_text,
        })
        return jsonify({"response": answer_text})
    except Exception as e:
        return jsonify({"response": f"Error generating content: {str(e)}"})


@app.route("/download_chat")
def download_chat():
    username = session['username']
    history = conversation_history[username]

    document = Document()
    document.add_heading(f"Conversation History for {username}", 0)

    for item in history:
        document.add_paragraph(f"You: {item['question']}", style='Intense Quote')
        document.add_paragraph(f"Bot: {item['answer']}")
        document.add_paragraph()

    doc_io = BytesIO()
    document.save(doc_io)
    doc_io.seek(0)

    return send_file(
        doc_io,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"conversation_{username}.docx"
    )

# The '/contact' route has been removed.

if __name__ == "__main__":
    app.run(debug=True)
